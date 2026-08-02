from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "公众号文章.md"
OUTPUT = ROOT / "公众号投稿稿-交易体系复盘.docx"

FONT = "Microsoft YaHei"
INK = "1F252B"
MUTED = "667079"
RED = "D64541"
LIGHT_RED = "FCEBEA"


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: float, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_inline_runs(paragraph, text: str, *, size: float = 10.5, color: str = INK, italic: bool = False) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, RED, bold=True, italic=italic)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, color, italic=italic)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, 9, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def paragraph_left_border(paragraph, color: str, size: int = 18, space: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    borders.append(left)


def configure_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in (
        ("Heading 1", 16, RED, 18, 9),
        ("Heading 2", 13, RED, 12, 6),
    ):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)

    if "Lead Callout" not in [style.name for style in document.styles]:
        callout = document.styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = document.styles["Lead Callout"]
    callout.font.name = FONT
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    callout.font.size = Pt(11)
    callout.font.color.rgb = RGBColor.from_string(INK)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(11)
    callout.paragraph_format.left_indent = Inches(0.16)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.line_spacing = 1.3

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("公众号投稿稿  |  交易体系复盘"), 8.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return document


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.line_spacing = 1.08
    set_run_font(title.add_run("近一年只赚2.70%，\n我却终于搭出了自己的交易体系"), 25, INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(subtitle.add_run("一个普通投资者从追涨杀跌，到七套策略的八年复盘"), 12.5, RED, bold=True)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.LEFT
    metadata.paragraph_format.space_after = Pt(13)
    set_run_font(metadata.add_run("个人投资复盘  |  数据截止 2026-07-28  |  历史结果不代表未来表现"), 8.5, MUTED)


def build() -> None:
    document = configure_document()
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    add_title_block(document)
    skip_title = True
    first_image = True

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# ") and skip_title:
            skip_title = False
            continue
        if line.startswith("> ") and first_image:
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            image_path = ROOT / image_match.group(2)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(5 if first_image else 7)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(6.35 if first_image else 6.2))
            first_image = False
            continue
        if line.startswith("*") and line.endswith("*"):
            paragraph = document.add_paragraph(style="Caption")
            add_inline_runs(paragraph, line[1:-1], size=8.5, color=MUTED)
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style="Lead Callout")
            paragraph_shading(paragraph, LIGHT_RED)
            paragraph_left_border(paragraph, RED)
            add_inline_runs(paragraph, line[2:], size=10.5)
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline_runs(paragraph, line)

    document.core_properties.title = "近一年只赚2.70%，我却终于搭出了自己的交易体系"
    document.core_properties.subject = "公众号投稿稿 - 个人投资体系复盘"
    document.core_properties.author = ""
    document.core_properties.keywords = "投资体系, 量化策略, 交易复盘"
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build()
